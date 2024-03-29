import pandas as pd
from util import Util
from docx import Document
import base64
import requests
import os
from datetime import datetime
from tabulate import tabulate
import missingno as msn
import matplotlib.pyplot as plt
import logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s : %(asctime)s : %(message)s', datefmt='%d-%b-%y %H:%M:%S')

class  MeticulousStatsReporter:
    def __init__(self, github_raw_url, data_file_path):
        
        """
        This constructor runs all statistical information generated functions and will update the stats in stats dataframes.
            Step-1: First, it will download the data configuration json and data from remote GitHub repository.
            Step-2: Classifying the dataframe columns based on their dtypes.
            Step-3: Creating statistics datafrmes to showcase stat reports.
            Step-4: It will update the datatype, mean, median, mode, null, etc... values stats.
        Parameters:
            github_raw_url : Data configaration json file url from remote GitHub repository.
            data_file_path : Input data file url from remote GitHub repository  .
        Return Values:
            No return values
        """

        # Creating required internal objs.
        util_obj = Util()
        
        # Step-1:
        ## Downloading data configuration json from remote GitHub repository.
        self.data_json = util_obj.download_data_config(github_raw_url)
        logging.debug(f"The data configuration json has been downloaded successfully.")

        ## Loading input data file as pandas df from remote GitHub repository.
        required_columns = self.data_json['required_columns'] if 'required_columns' in self.data_json else None
        self.data = util_obj.load_input_data(data_file_path, required_columns)
        self.data.head()
        logging.debug(f"The input data loaded successfully.")

        # Step-2:
        ## Classifying the dataframe columns based on its dtypes.
        self.datatype_classified_clms = dict()
        util_obj.classifying_numeric_and_obj_clms(self)
        logging.debug(f"""The dataframe columns are classified successfully based on dtypes.
                     The classified datatype columns are {self.datatype_classified_clms}""")
        
        # Step-3:
        ## Creating statistics datafrmes to showcase stat reports.
        self.statistics_dfs_dict = dict()
        self.create_statistics_dfs()

        # Step-4:
        self.update_datatype_stats() # Updating datatype stats
        self.update_mean_median_mode() # Updating mean, median, mode values
        self.update_null_values_count() # Updating null values
        fname = github_raw_url.split('/')[-1]
        self.stats_report(fname)

        # self.logs = [f"File Name: {fname}                 Date:", "rama"]
        # self.save_logs_to_doc(self.logs)

        # # Displaying final statistics dataframes.
        # for key, df in self.statistics_dfs_dict.items():
        #     print(df)

    def create_statistics_dfs(self):
        for key, clm_list in self.datatype_classified_clms.items():
            if len(clm_list) > 0:
                self.statistics_dfs_dict[key] = pd.DataFrame(index=clm_list ,
                                                             columns=['Data_Types', 'mean', 'median', 'mode', 'Null_Values_Count'])

    def update_datatype_stats(self):
        for key in self.statistics_dfs_dict:
            clm_datatype_dict = self.data[self.datatype_classified_clms[key]].dtypes
            self.statistics_dfs_dict[key]['Data_Types'] = self.statistics_dfs_dict[key].index.map(clm_datatype_dict)

    def update_mean_median_mode(self):
        mean_median_mode_values = dict()
        if 'numeric_columns' in self.statistics_dfs_dict:
            mean_median_mode_values['mean'] = {numaric_clm: self.data[numaric_clm].mean() for numaric_clm in self.datatype_classified_clms['numeric_columns']}
            mean_median_mode_values['median'] = {numaric_clm: self.data[numaric_clm].median() for numaric_clm in self.datatype_classified_clms['numeric_columns']}
            mean_median_mode_values['mode'] = {numaric_clm: self.data[numaric_clm].mode().iloc[0] for numaric_clm in self.datatype_classified_clms['numeric_columns']}
            
            for key, val in mean_median_mode_values.items():
                self.statistics_dfs_dict['numeric_columns'][key] = self.statistics_dfs_dict['numeric_columns'].index.map(val)

    def update_null_values_count(self):
        # Updating null values count in stats df from all types of columns with for loop.
        for key in self.statistics_dfs_dict:
            null_values_count_dict = {numaric_clm: self.data[numaric_clm].isna().sum() for numaric_clm in self.datatype_classified_clms[key]}
            self.statistics_dfs_dict[key]['Null_Values_Count'] = self.statistics_dfs_dict[key].index.map(null_values_count_dict)

    def save_logs_to_doc(self, logs, doc_filename='logs_document.docx'):
        # Use the GitHub Actions workspace directory
        # workspace_dir = os.environ.get('GITHUB_WORKSPACE', '_')
        # doc_path = os.path.join(workspace_dir, doc_filename)
        document = Document()

        # Add a title to the document
        document.add_heading('Generated Logs', level=1)

        # Add logs to the document
        for log in logs:
            document.add_paragraph(log)

        # Save the document
        document.save(doc_filename)
        print(f"Logs saved to {doc_filename}")

    def stats_report(self, fname):
        stats_report_str = f"""
                          STATISTICAL REPORT
Report ID: 12345                           Date: {datetime.now()}
File Name: {fname}

=============================================================================="""
        print(stats_report_str)
        for key, df in self.statistics_dfs_dict.items():
            print(f"\nThe {key} reports :\n")
            print(tabulate(df, headers = 'keys', tablefmt = 'psql'))
            print("\nThe null values graphs: \n")
            msn.matrix(self.data)
            plt.savefig('matrix_plot.png')  # Save the matrix plot as an image file
            plt.clf()  # Clear the figure for the next plot
            plt.show()